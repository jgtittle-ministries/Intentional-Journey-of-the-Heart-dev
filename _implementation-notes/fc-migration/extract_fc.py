import docx, os, zipfile
import xml.etree.ElementTree as ET
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = r"C:\Users\jgtit\OneDrive\Documents\Intentional Journey of the Heart\Reference Articles for pdf in GitHub\Formation Companion.docx"
OUT = r"C:\Users\jgtit\claude\_work\Intentional-Journey-of-the-Heart-dev\_implementation-notes\fc-migration\fc-extracted.md"

d = docx.Document(SRC)

def style_prefix(stylename):
    s = (stylename or "").lower()
    if s == "title":          return "# "
    if s.startswith("heading 1"): return "# "
    if s.startswith("heading 2"): return "## "
    if s.startswith("heading 3"): return "### "
    if s.startswith("list") or s.startswith("bullet"): return "- "
    return ""

out = []
for child in d.element.body.iterchildren():
    if child.tag == qn('w:p'):
        para = Paragraph(child, d)
        t = para.text.strip()
        if not t:
            out.append("")
            continue
        out.append(style_prefix(para.style.name if para.style else "") + t)
    elif child.tag == qn('w:tbl'):
        tbl = Table(child, d)
        out.append("")
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                out.append("|" + "|".join(["---"] * len(cells)) + "|")
        out.append("")

text = "\n".join(out)

# footnotes
fn = []
z = zipfile.ZipFile(SRC)
if "word/footnotes.xml" in z.namelist():
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = ET.fromstring(z.read("word/footnotes.xml"))
    for fnote in root.findall("w:footnote", ns):
        fid = fnote.get(qn('w:id'))
        txt = "".join(t.text or "" for t in fnote.iter(qn('w:t')))
        if txt.strip() and fid not in ("-1", "0"):
            fn.append("[^%s]: %s" % (fid, txt.strip()))

os.makedirs(os.path.dirname(OUT), exist_ok=True)
with open(OUT, "w", encoding="utf-8") as f:
    f.write(text)
    if fn:
        f.write("\n\n## Footnotes (extracted)\n\n" + "\n".join(fn))

print("wrote", OUT)
print("chars:", len(text), "footnotes:", len(fn), "paras:", len(d.paragraphs), "tables:", len(d.tables))
