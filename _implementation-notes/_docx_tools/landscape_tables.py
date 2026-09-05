# Lay a pandoc-built docx out in landscape with narrow margins and small table text.
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT

path = sys.argv[1]
doc = Document(path)
for s in doc.sections:
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Inches(11), Inches(8.5)
    s.left_margin = s.right_margin = Inches(0.5)
    s.top_margin = s.bottom_margin = Inches(0.6)
for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
    try:
        st = doc.styles[name]
        st.font.size = Pt(9.5)
        st.paragraph_format.space_after = Pt(3)
    except KeyError:
        pass
for t in doc.tables:
    t.autofit = True
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(8)
doc.save(path)
print("landscape", path)
